"""
Exporter Module.
Provides functions to export markdown summaries and notes to PDF (via reportlab) and DOCX (via python-docx).
"""

import os
import re
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document


def markdown_to_html_like(text: str) -> str:
    """Convert standard markdown tags to reportlab Paragraph compatible XML tags."""
    # Escape standard XML characters
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    # Restore tags we want to format
    text = text.replace("&lt;b&gt;", "<b>").replace("&lt;/b&gt;", "</b>")
    text = text.replace("&lt;i&gt;", "<i>").replace("&lt;/i&gt;", "</i>")
    
    # Match markdown **bold** to <b>bold</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Match markdown *italic* to <i>italic</i>
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    # Convert `code` to typewriter font
    text = re.sub(r'`(.*?)`', r'<font face="Courier">\1</font>', text)
    return text


def generate_pdf(markdown_text: str, filename: str) -> str:
    """Generate a styled PDF from markdown content."""
    doc = SimpleDocTemplate(
        filename,
        pagesize=letter,
        rightMargin=45,
        leftMargin=45,
        topMargin=45,
        bottomMargin=45
    )
    styles = getSampleStyleSheet()
    
    # Premium typography settings
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor='#0f172a',
        spaceAfter=15,
        alignment=TA_CENTER
    )
    
    h1_style = ParagraphStyle(
        'H1',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor='#1e3a8a',
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor='#334155',
        spaceAfter=5
    )
    
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )
    
    story = []
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 3))
            continue
            
        # Parse headers
        if line.startswith('# '):
            text = markdown_to_html_like(line[2:])
            story.append(Paragraph(text, title_style))
            story.append(Spacer(1, 8))
        elif line.startswith('## '):
            text = markdown_to_html_like(line[3:])
            story.append(Paragraph(text, h1_style))
        elif line.startswith('### '):
            text = markdown_to_html_like(line[4:])
            story.append(Paragraph(text, h1_style))
        elif line.startswith('- ') or line.startswith('* '):
            text = markdown_to_html_like(line[2:])
            story.append(Paragraph(f"&bull; {text}", bullet_style))
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or re.match(r'^\d+\.\s', line):
            text = markdown_to_html_like(line)
            story.append(Paragraph(text, bullet_style))
        else:
            text = markdown_to_html_like(line)
            story.append(Paragraph(text, body_style))
            
    doc.build(story)
    return filename


def generate_docx(markdown_text: str, filename: str) -> str:
    """Generate a cleanly formatted DOCX file from markdown content."""
    doc = Document()
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            # Simple bold formatting parser for document run
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
                    
    doc.save(filename)
    return filename
